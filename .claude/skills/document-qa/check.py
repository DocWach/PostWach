#!/usr/bin/env python3
"""document-qa /check -- pre-signoff consistency & completeness scan for a .docx.

Catches the "simple things" that eyeballing misses: duplicate/missing figure & table
labels, captions orphaned in text boxes, missing captions/callouts, inconsistent table
and caption formatting, leftover scaffolding, and heading-hierarchy outliers.

Usage:  python check.py <path-to.docx>
Stdlib + python-docx only. Pair with the render-and-look gate (see SKILL.md).
"""
import sys, re, os
from collections import Counter, defaultdict
import docx
from docx.oxml.ns import qn

FAILS = []; WARNS = []; INFOS = []
def fail(m): FAILS.append(m)
def warn(m): WARNS.append(m)
def info(m): INFOS.append(m)

def ptext(p):
    return "".join(t.text or "" for t in p.iter(qn('w:t')))

def in_textbox(p):
    el = p.getparent()
    while el is not None:
        if el.tag.split('}')[-1] == 'txbxContent':
            return True
        el = el.getparent()
    return False

def main(path):
    d = docx.Document(path)
    allp = list(d.element.iter(qn('w:p')))            # includes text boxes
    body_text = "\n".join(p.text for p in d.paragraphs)

    # ---- captions ----
    cap_re = re.compile(r"^(Figure|Table)\s+(\d+)\s*\.")   # a caption = "Figure N." / "Table N."
    caps = []   # (kind, num, in_textbox, text)
    for p in allp:
        t = ptext(p).strip()
        m = cap_re.match(t)
        if m and len(t) < 300:
            caps.append((m.group(1), int(m.group(2)), in_textbox(p), t[:60]))
    for kind in ("Figure", "Table"):
        nums = [c[1] for c in caps if c[0] == kind]
        info(f"{kind} captions found: {sorted(nums)}")
        for n, c in Counter(nums).items():
            if c > 1:
                fail(f"DUPLICATE {kind} {n} caption appears {c}x "
                     f"(check for an old copy in a text box)")
        if nums:
            expected = set(range(1, max(nums) + 1))
            missing = expected - set(nums)
            if missing:
                fail(f"{kind} numbering gap: missing {sorted(missing)}")
        for kd, n, tb, txt in caps:
            if kd == kind and tb:
                fail(f"{kind} {n} caption is inside a TEXT BOX (orphan/floating): {txt!r}")

    # ---- in-text callouts (referenced by number in prose) ----
    for kind in ("Figure", "Table"):
        for n in sorted({c[1] for c in caps if c[0] == kind}):
            # the caption contributes one "Kind N" occurrence; a real callout adds more
            refs = len(re.findall(rf"\b{kind}\s+{n}\b", body_text))
            if refs <= 1:
                warn(f"{kind} {n} has no in-text callout (only its caption) — "
                     f"reference it by number in the prose")

    # ---- tables: formatting consistency ----
    def hfill(c):
        tc = c._tc.find(qn('w:tcPr'))
        if tc is not None:
            s = tc.find(qn('w:shd'))
            if s is not None: return s.get(qn('w:fill'))
        return None
    def col0_bold(t):
        return any(any(x.bold for x in r.cells[0].paragraphs[0].runs)
                   for r in t.rows[1:] if r.cells[0].paragraphs[0].runs)
    def cell_sizes(t):
        s = set()
        for r in t.rows[1:]:
            for run in r.cells[0].paragraphs[0].runs:
                if run.text.strip(): s.add(run.font.size)
        return s
    tabs = d.tables
    info(f"tables found: {len(tabs)}")
    fills = [hfill(t.rows[0].cells[0]) for t in tabs]
    borders = [t._tbl.tblPr.find(qn('w:tblBorders')) is not None for t in tabs]
    bolds = [col0_bold(t) for t in tabs]
    if len(set(fills)) > 1:
        warn(f"table header shading inconsistent across tables: {fills}")
    if len(set(borders)) > 1:
        fail(f"table border style inconsistent (some gridded, some not): {borders}")
    if len(set(bolds)) > 1:
        warn(f"first-column bold inconsistent across tables (row-label style on some, not others): {bolds}")
    def halign(t):
        return frozenset(str(p.alignment) for r in t.rows[1:] for c in r.cells for p in c.paragraphs)
    def valign(t):
        out = set()
        for r in t.rows:
            for c in r.cells:
                tc = c._tc.find(qn('w:tcPr'))
                va = tc.find(qn('w:vAlign')) if tc is not None else None
                out.add(va.get(qn('w:val')) if va is not None else None)
        return frozenset(out)
    if len({halign(t) for t in tabs}) > 1:
        warn(f"table cell horizontal-alignment inconsistent across tables: "
             f"{[sorted(map(str, halign(t))) for t in tabs]}")
    if len({valign(t) for t in tabs}) > 1:
        warn(f"table cell vertical-alignment inconsistent across tables (some centered, some top): "
             f"{[sorted(map(str, valign(t))) for t in tabs]}")
    # duplicate tables (same header signature)
    sigs = ["|".join(c.text.strip().lower()[:15] for c in t.rows[0].cells) for t in tabs]
    for sig, c in Counter(sigs).items():
        if c > 1 and sig.strip("|"):
            warn(f"possible duplicate tables sharing header {sig!r} ({c}x)")

    # ---- caption style consistency ----
    styles = []
    for p in d.paragraphs:
        t = p.text.strip()
        if cap_re.match(t):
            prefix_bold = bool(p.runs) and bool(p.runs[0].bold)
            title_bold = any(r.bold for r in p.runs[1:] if r.text.strip())
            sizes = tuple(sorted({str(r.font.size) for r in p.runs if r.text.strip()}))
            styles.append((t[:22], prefix_bold, title_bold, sizes))
    patt = Counter((s[1], s[2]) for s in styles)
    if len(patt) > 1:
        warn("caption bold pattern inconsistent (some prefix-only-bold, some all-bold): "
             + "; ".join(f"{s[0]!r}:prefixbold={s[1]},titlebold={s[2]}" for s in styles))
    allsizes = set(sz for s in styles for sz in s[3])
    if len(allsizes) > 1:
        warn(f"caption font sizes inconsistent: {allsizes}")

    # ---- scaffolding / placeholders ----
    scaff = re.findall(r"\[(?:Define|cite|cited|PLACEHOLDER|TODO|TBD|INSERT)[^\]]*\]"
                       r"|\(include Subcontractors[^)]*\)|\bTBD\b|\bTODO\b", body_text)
    for s in set(scaff):
        fail(f"scaffolding/placeholder left in document: {s!r}")

    # ---- heading hierarchy (top-level 'N. Title') ----
    hs = []
    for p in d.paragraphs:
        t = p.text.strip()
        if re.match(r"^\d{1,2}\.\s+[A-Z]", t) and len(t) < 140 and p.runs:
            hs.append((t[:30], p.runs[0].bold, str(p.runs[0].font.size)))
    if hs:
        if len({h[1] for h in hs}) > 1:
            fail("top-level section headers: some bold, some not: "
                 + ", ".join(f"{h[0]!r}={h[1]}" for h in hs if not h[1]))
        if len({h[2] for h in hs}) > 1:
            warn(f"top-level section header sizes inconsistent: {Counter(h[2] for h in hs)}")

    # ---- report ----
    print(f"\n=== document-qa scan: {os.path.basename(path)} ===")
    for m in INFOS: print(f"  INFO  {m}")
    for m in WARNS: print(f"  WARN  {m}")
    for m in FAILS: print(f"  FAIL  {m}")
    print(f"\n  {len(FAILS)} FAIL / {len(WARNS)} WARN")
    print("  Reminder: a clean scan is necessary, not sufficient. Now RENDER and LOOK.")
    return 1 if FAILS else 0

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("usage: python check.py <path-to.docx>"); sys.exit(2)
    sys.exit(main(sys.argv[1]))
