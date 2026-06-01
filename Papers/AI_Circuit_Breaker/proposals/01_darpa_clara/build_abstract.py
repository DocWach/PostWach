"""Populate the CLARA Abstract Template .docx with our content."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree
from pathlib import Path

BASE = Path(__file__).parent
TEMPLATE = BASE / "Compositional+Learning-And-Reasoning+for+AI+Complex+Systems+Engineering+(CLARA)" / "CLARA_-_Abstract_Template.docx"
OUTPUT = BASE / "clara_abstract_completed.docx"

doc = Document(str(TEMPLATE))

# ============================================================
# PAGE 1: COVER SHEET
# ============================================================
table = doc.tables[0]

# Row 0: Title
table.rows[0].cells[1].text = "Morphism-Grounded Compositional Assurance for Autonomous AI Systems"
for p in table.rows[0].cells[1].paragraphs:
    for r in p.runs:
        r.font.size = Pt(11)

# Row 1: Organization
table.rows[1].cells[1].text = "University of Arizona"
for p in table.rows[1].cells[1].paragraphs:
    for r in p.runs:
        r.font.size = Pt(11)

# Row 2: Technical POC
cell = table.rows[2].cells[1]
cell.text = ""
lines = [
    "Name: Paul F. Wach, Ph.D.",
    "Mailing Address: Dept. of Systems & Industrial Engineering, University of Arizona, Tucson, AZ 85721",
    "Telephone: [PHONE]",
    "Email: [EMAIL]",
]
for i, line in enumerate(lines):
    p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
    run = p.add_run(line)
    run.font.size = Pt(11)

# Row 3: Administrative POC
cell = table.rows[3].cells[1]
cell.text = ""
lines = [
    "Name: [UA Sponsored Projects Contact]",
    "Mailing Address: [UA Sponsored Projects Address]",
    "Telephone: [PHONE]",
    "Email: [EMAIL]",
]
for i, line in enumerate(lines):
    p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
    run = p.add_run(line)
    run.font.size = Pt(11)

# Replace logo placeholder
for para in doc.paragraphs:
    if "[PRIME ORGANIZATION LOGO]" in para.text:
        para.text = "[University of Arizona Logo]"
        for r in para.runs:
            r.font.size = Pt(11)
        break

# ============================================================
# PAGE 2: TECHNICAL DESCRIPTION
# ============================================================

# The technical content, split into (heading_bool, text) pairs
content = [
    (True,  "1. What is the proposed work attempting to accomplish?"),
    (False, "We propose a compositional ML/AR framework that provides formal, composable "
            "assurance guarantees for AI systems of systems. The framework couples an AR "
            "scaffold \u2014 OWL 2 DL ontology with SHACL constraint shapes and logic-based "
            "graduated response rules \u2014 with ML components \u2014 transformer-based neural "
            "network encodings and Bayesian statistical process control \u2014 to continuously "
            "measure and bound the trustworthiness of composed AI subsystems. AR defines the "
            "compositional structure and enforces validity constraints on ML outputs; ML "
            "provides runtime evidence that AR reasons over. The composition is inherent in "
            "the measurement theory, not bolted on. Phase 1 delivers open-source software "
            "(Apache 2.0) composing Logic Programs (AR) with Neural Networks (ML), evaluated "
            "on a safety-critical task domain with explicit SOA benchmarks."),
    (True,  "2. How is the work performed today, and what are the limitations?"),
    (False, "Current approaches bolt AR onto ML as external guardrails \u2014 the failure mode "
            "CLARA identifies. Three limitations persist. (1) No formal composition theory: "
            "trust in a system of systems is assumed, not derived from component trust; no "
            "framework bounds composite trustworthiness. (2) No measurement rigor: runtime "
            "monitors use ad-hoc thresholds with no traceability to specifications and no "
            "uncertainty quantification. (3) Single-axis monitoring: existing approaches check "
            "structural fidelity or behavioral accuracy but never both, missing half the "
            "failure space. DeepProbLog and similar systems begin to address AR-based ML but "
            "lack polynomial tractability and high-expressiveness AR features (defeasible "
            "argumentation, restraint) that CLARA requires."),
    (True,  "3. What is new in your approach, and why will it succeed?"),
    (False, "Core innovation. We ground AI trustworthiness in systems-theoretic morphisms "
            "\u2014 structure-preserving mappings between formal system models. An AI agent\u2019s "
            "internal model Z_ai and the physical world Z_real are formalized as system "
            "specifications. Trustworthiness is the measurable quality of the mapping "
            "h: Z_ai \u2192 Z_real, decomposed along two orthogonal axes: structural quality "
            "\u03c3 (degree of homomorphism \u2014 an AR computation via ontology alignment and "
            "graph homomorphism checking) and behavioral quality D (output distance \u2014 an "
            "ML computation via embedding similarity and sensor fusion). Neither axis alone "
            "suffices; both are jointly necessary, making the ML/AR composition inherent "
            "rather than tacked on."),
    (False, "Compositional guarantees. Morphism quality bounds compose functorially: "
            "\u03c3_total \u2264 min(\u03c3_component) and D_total \u2264 \u03a3(D_component), "
            "directly addressing CLARA\u2019s hierarchical composition requirement. SHACL "
            "validation ensures ontological consistency across composed subsystems. PROV-O "
            "provenance graphs provide hierarchical, fine-grained explainability with low "
            "unfolding depth."),
    (False, "Why it will meet CLARA\u2019s metrics. (1) Verifiability without loss of "
            "performance: morphism bounds provide automatic proofs of composite trust; SOA "
            "benchmarks will be established against ad-hoc threshold monitors. "
            "(2) Multiplicity of AI kinds: the framework is kind-agnostic \u2014 Phase 1 "
            "composes Neural Networks (ML) with Logic Programs (AR); Phase 2 extends to "
            "additional kinds. (3) Polynomial time complexity: structural morphism quality "
            "(\u03c3) computes in O(n\u00b7k) via approximate nearest-neighbor search; "
            "behavioral quality (D) is O(m) sensor fusion. (4) Composed task reliability "
            "> SOA: dual-axis monitoring detects failures that single-axis monitors miss, "
            "demonstrated via hallucination injection at 1\u201325% rates on the selected "
            "task domain."),
    (False, "Technical basis. The PI has published directly on morphism-based system "
            "verification, degree-of-homomorphism metrics, conjoining systems-theoretic and "
            "DEVS formalisms, and pairing Bayesian methods with systems theory for T&E of "
            "learning-based systems \u2014 the precise theoretical foundations this approach "
            "requires."),
]

# Find the instruction paragraph (para 18) and list items (20-22)
# Remove instruction text + list items, replace with our content
body = doc.element.body

# Identify elements to remove: paragraphs 18-23 (instruction, blank, 3 list items, trailing blank)
elems_to_remove = []
for i in range(18, min(24, len(doc.paragraphs))):
    elems_to_remove.append(doc.paragraphs[i]._element)

# Find insertion point (just before paragraph 18)
insert_ref = elems_to_remove[0]

# Remove old content
for elem in elems_to_remove:
    body.remove(elem)

# Helper to create a paragraph element with styled text
def make_paragraph(text, bold=False, font_size=20):
    """Create a w:p element. font_size in half-points (20 = 10pt)."""
    p = etree.SubElement(body, qn("w:p"))

    # For bold-headed paragraphs where first sentence is bold
    if not bold:
        # Check if text starts with a bold-leader pattern like "Core innovation."
        bold_leaders = [
            "Core innovation.", "Compositional guarantees.",
            "Why it will meet CLARA\u2019s metrics.", "Technical basis."
        ]
        leader = None
        for bl in bold_leaders:
            if text.startswith(bl):
                leader = bl
                break

        if leader:
            # Bold run for leader
            run1 = etree.SubElement(p, qn("w:r"))
            rpr1 = etree.SubElement(run1, qn("w:rPr"))
            etree.SubElement(rpr1, qn("w:b"))
            sz1 = etree.SubElement(rpr1, qn("w:sz"))
            sz1.set(qn("w:val"), str(font_size))
            szCs1 = etree.SubElement(rpr1, qn("w:szCs"))
            szCs1.set(qn("w:val"), str(font_size))
            t1 = etree.SubElement(run1, qn("w:t"))
            t1.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t1.text = leader

            # Normal run for rest
            rest = text[len(leader):]
            if rest:
                run2 = etree.SubElement(p, qn("w:r"))
                rpr2 = etree.SubElement(run2, qn("w:rPr"))
                sz2 = etree.SubElement(rpr2, qn("w:sz"))
                sz2.set(qn("w:val"), str(font_size))
                szCs2 = etree.SubElement(rpr2, qn("w:szCs"))
                szCs2.set(qn("w:val"), str(font_size))
                t2 = etree.SubElement(run2, qn("w:t"))
                t2.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                t2.text = rest
            return p

    # Simple case: one run
    run = etree.SubElement(p, qn("w:r"))
    rpr = etree.SubElement(run, qn("w:rPr"))
    sz = etree.SubElement(rpr, qn("w:sz"))
    sz.set(qn("w:val"), str(font_size))
    szCs = etree.SubElement(rpr, qn("w:szCs"))
    szCs.set(qn("w:val"), str(font_size))
    if bold:
        etree.SubElement(rpr, qn("w:b"))
    t = etree.SubElement(run, qn("w:t"))
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return p

# Insert our content paragraphs in order
# We need to insert before the element that was after paragraph 23
# Since we removed those elements, insert_ref is gone. Let's find the right spot.
# The "Abstract Technical Description" heading is paragraph 16.
# After removal, the next element after para 17 (blank after heading) is wherever we need to insert.
# Let's find paragraph 16's element and insert after paragraph 17.

heading_elem = doc.paragraphs[16]._element  # "Abstract Technical Description"
# Para 17 is a blank line after heading
blank_after_heading = doc.paragraphs[17]._element

# Insert content after the blank line
prev = blank_after_heading
for is_heading, text in content:
    p_elem = make_paragraph(text, bold=is_heading, font_size=20)
    body.remove(p_elem)  # remove from end of body where SubElement put it
    prev.addnext(p_elem)
    prev = p_elem

# Save
doc.save(str(OUTPUT))
print(f"Saved completed abstract to: {OUTPUT.name}")
print("Done. Fill in [PHONE], [EMAIL], Admin POC, and UA logo before submission.")
